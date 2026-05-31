import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from typing import Dict, Any

class ExportService:
    @staticmethod
    def generate_ieee_docx(data: Dict[str, Any]) -> io.BytesIO:
        doc = docx.Document()
        
        # IEEE Style - Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(data.get("title", "Untitled Research Paper"))
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Abstract
        doc.add_heading("Abstract", level=1)
        doc.add_paragraph(data.get("abstract", ""))
        
        # Keywords
        keywords_para = doc.add_paragraph()
        keywords_para.add_run("Keywords—").bold = True
        keywords_para.add_run(", ".join(data.get("keywords", [])))
        
        # Sections
        for section in data.get("sections", []):
            doc.add_heading(section.get("title", ""), level=1)
            doc.add_paragraph(section.get("content", ""))
            
        # References
        doc.add_heading("References", level=1)
        for ref in data.get("references", []):
            doc.add_paragraph(ref, style='List Number')
            
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def generate_ieee_latex(data: Dict[str, Any]) -> str:
        # Simplified LaTeX template
        sections_latex = ""
        for section in data.get("sections", []):
            sections_latex += f"\\section{{{section.get('title', '')}}}\n{section.get('content', '')}\n\n"
            
        references_latex = ""
        for i, ref in enumerate(data.get("references", [])):
            references_latex += f"\\bibitem{{ref{i+1}}} {ref}\n"
            
        latex_template = f"""
\\documentclass[conference]{{IEEEtran}}
\\begin{{document}}
\\title{{{data.get('title', '')}}}
\\author{{\\IEEEauthorblockN{{Author Name}} \\IEEEauthorblockA{{Department, University}}}}
\\maketitle

\\begin{{abstract}}
{data.get('abstract', '')}
\\end{{abstract}}

\\begin{{IEEEkeywords}}
{', '.join(data.get('keywords', []))}
\\end{{IEEEkeywords}}

{sections_latex}

\\begin{{thebibliography}}{{1}}
{references_latex}
\\end{{thebibliography}}
\\end{{document}}
        """
        return latex_template
